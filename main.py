import os
import re
import shutil
import tempfile
from typing import List, Dict, Any, Tuple, Optional
import json
import asyncio
from datetime import datetime, timedelta # Import timedelta
import logging # Import logging

# FastAPI imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Depends, status
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse # Import HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field # Still used for request body validation

# SQLAlchemy imports for database ORM
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, ForeignKey, JSON
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship
from sqlalchemy.sql import func

# Password hashing
from passlib.context import CryptContext

# JWT for authentication
from jose import JWTError, jwt

# LangChain imports
from langchain_openai import ChatOpenAI # Using ChatOpenAI for OpenRouter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.output_parsers import JsonOutputParser, PydanticOutputParser # CORRECTED: Import PydanticOutputParser
from pydantic import BaseModel as LangchainBaseModel, Field as LangchainField # Use pydantic for Langchain models
from langchain.text_splitter import RecursiveCharacterTextSplitter, Language # Import Language for code splitting
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders.parsers import LanguageParser # Still used, but implicitly by RecursiveCharacterTextSplitter.from_language
from langchain.document_loaders.git import GitLoader # Keep GitLoader
from langchain_core.documents import Document

# For resume parsing
import PyPDF2
import docx

# For GitHub content fetching (if README.md is preferred over cloning entire repo for some cases)
import requests

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

# --- Logging Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Disable tokenizers parallelism warning (often seen with multiprocessing/Flask reloader)
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# --- Configuration ---
# Database Configuration
DATABASE_URL = os.getenv("DATABASE_URL")
logger.info(f"DATABASE_URL loaded: {DATABASE_URL}")

# JWT Configuration
SECRET_KEY = os.getenv("SECRET_KEY")
if not SECRET_KEY:
    logger.warning("SECRET_KEY not set in .env! Using a default, which is INSECURE for production.")
    SECRET_KEY = "insecure-default-secret-key-please-change-this" # Fallback for development
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 300 # 5 hours

# LLM and Embeddings Configuration
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "") # Default to empty if not set

# Initialize LLM (Using ChatOpenAI for OpenRouter)
llm = ChatOpenAI(
    model_name="deepseek/deepseek-r1-0528-qwen3-8b:free", # Specific model as requested
    openai_api_key=OPENROUTER_API_KEY, # Use the API key from .env
    openai_api_base="https://openrouter.ai/api/v1", # OpenRouter API base URL
    temperature=0.7 # Keep temperature as before
)
logger.info("LangChain LLM (ChatOpenAI with OpenRouter/DeepSeek) initialized.")

# Initialize Embeddings (HuggingFace sentence-transformers/all-MiniLM-L6-v2)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
logger.info("HuggingFace Embeddings model loaded.")

# --- Database Setup ---
Base = declarative_base()
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

# Create database tables on startup
try:
    Base.metadata.create_all(bind=engine)
    logger.info("Database tables checked/created successfully.")
except Exception as e:
    logger.error(f"Error creating database tables: {e}")
    # Depending on severity, you might want to exit or raise here

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Password hashing context
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# OAuth2 scheme for authentication
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# --- Database Models ---
class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(255), unique=True, index=True, nullable=False)
    hashed_password = Column(String(255), nullable=False)

    sessions = relationship("InterviewSession", back_populates="user")

class InterviewSession(Base):
    __tablename__ = "interview_sessions"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    start_time = Column(DateTime, default=func.now(), nullable=False)
    end_time = Column(DateTime, nullable=True)
    resume_text_snippet = Column(Text, nullable=True) # Store first N chars of resume
    github_knowledge_summary = Column(Text, nullable=True) # Store summary of processed github repos

    user = relationship("User", back_populates="sessions")
    conversations = relationship("InterviewConversation", back_populates="session", cascade="all, delete-orphan")
    feedback = relationship("InterviewFeedback", back_populates="session", uselist=False, cascade="all, delete-orphan")

class InterviewConversation(Base):
    __tablename__ = "interview_conversations"
    id = Column(Integer, primary_key=True, index=True)
    interview_session_id = Column(Integer, ForeignKey("interview_sessions.id"), nullable=False)
    role = Column(String(50), nullable=False) # 'ai' or 'user'
    text = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=func.now(), nullable=False)

    session = relationship("InterviewSession", back_populates="conversations")

class InterviewFeedback(Base):
    __tablename__ = "interview_feedback"
    id = Column(Integer, primary_key=True, index=True)
    interview_session_id = Column(Integer, ForeignKey("interview_sessions.id"), unique=True, nullable=False)
    technical_rating = Column(Integer, nullable=True)
    technical_tips = Column(Text, nullable=True) # Stored as JSON string
    hr_rating = Column(Integer, nullable=True)
    hr_tips = Column(Text, nullable=True) # Stored as JSON string

    session = relationship("InterviewSession", back_populates="feedback")

# Create database tables
Base.metadata.create_all(bind=engine)

# --- Pydantic Models for Structured Output (LangChain) ---
class TechnicalFeedback(LangchainBaseModel):
    """Feedback from the technical agent."""
    technical_knowledge_rating: int = LangchainField(description="Rating of technical knowledge out of 5.")
    technical_tips: List[str] = LangchainField(description="List of actionable tips to improve technical answers.")

class HRFeedback(LangchainBaseModel):
    """Feedback from the HR agent."""
    communication_skills_rating: int = LangchainField(description="Rating of communication skills out of 5.")
    communication_tips: List[str] = LangchainField(description="List of actionable tips to improve communication and soft skills.")

# NEW: Pydantic Model for a single interview question
class InterviewQuestion(LangchainBaseModel):
    question: str = LangchainField(description="A single, direct interview question for the candidate.")

# NEW: Output parser for the interview question
question_parser = PydanticOutputParser(pydantic_object=InterviewQuestion)
question_format_instructions = question_parser.get_format_instructions()


# --- FastAPI App Setup ---
app = FastAPI()

# Allow CORS for frontend interaction (important for local development)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Adjust in production to specific frontend origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files to serve other static assets (e.g., /static/images/logo.png)
app.mount("/static", StaticFiles(directory="static"), name="static_assets")
logger.info("Static files mounted at /static.")

# Global dictionary for in-memory active interview session data
# This stores temporary data for an ongoing interview, not persisted until finished
active_sessions: Dict[str, Dict[str, Any]] = {}
logger.info("Active sessions dictionary initialized.")

# --- Authentication Utilities ---
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    logger.info(f"Access token created for sub: {data.get('sub')}")
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            logger.warning("JWT payload missing 'sub' (username).")
            raise credentials_exception
    except JWTError as e:
        logger.error(f"JWT decoding error: {e}")
        raise credentials_exception
    user = db.query(User).filter(User.username == username).first()
    if user is None:
        logger.warning(f"User '{username}' not found in DB during token validation.")
        raise credentials_exception
    logger.info(f"User '{username}' authenticated successfully.")
    return user

# --- Resume Processing Functions ---

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
        logger.info(f"Text extracted from PDF: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        logger.info(f"Text extracted from DOCX: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}")
        return ""

def extract_github_links(text: str) -> List[str]:
    """Extracts GitHub repository URLs from text."""
    github_patterns = [
        r'https?://github\.com/([a-zA-Z0-9_-]+/[a-zA-Z0-9_.-]+)(?:/tree/[a-zA-Z0-9_.-]+)?(?:/blob/[a-zA-Z0-9_.-]+)?',
        r'git@github\.com:([a-zA-Z0-9_-]+/[a-zA-Z0-9_.-]+)\.git'
    ]
    links = []
    for pattern in github_patterns:
        found_repos = re.findall(pattern, text)
        for repo_path in found_repos:
            links.append(f"https://github.com/{repo_path.replace('.git', '')}")
    unique_links = list(set(links))
    logger.info(f"Extracted GitHub links: {unique_links}")
    return unique_links

async def clone_and_embed_github_repo(repo_url: str, session_id: str) -> List[Document]:
    """
    Clones a GitHub repository and processes its code files for embedding.
    Returns a list of LangChain Documents from the cloned code.
    """
    temp_dir = tempfile.mkdtemp(prefix=f"github_clone_{session_id}_")
    logger.info(f"Cloning {repo_url} into {temp_dir}")
    try:
        loader = GitLoader(repo_url=repo_url, branch="main", clone_dir=temp_dir)
        all_docs = loader.load()
        logger.info(f"Loaded {len(all_docs)} documents from {repo_url}")

        processed_docs = []
        
        # Define language-specific splitters
        python_splitter = RecursiveCharacterTextSplitter.from_language(
            language=Language.PYTHON, chunk_size=1000, chunk_overlap=200
        )
        js_splitter = RecursiveCharacterTextSplitter.from_language(
            language=Language.JS, chunk_size=1000, chunk_overlap=200
        )
        markdown_splitter = RecursiveCharacterTextSplitter.from_language(
            language=Language.MARKDOWN, chunk_size=1000, chunk_overlap=200
        )
        generic_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

        for doc in all_docs:
            file_path = doc.metadata.get('file_path', '')
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Determine which splitter to use based on file extension
            if file_extension == '.py':
                chunks = python_splitter.split_documents([doc])
                logger.debug(f"Split {file_path} as Python.")
            elif file_extension == '.js' or file_extension == '.ts':
                chunks = js_splitter.split_documents([doc])
                logger.debug(f"Split {file_path} as JavaScript/TypeScript.")
            elif file_extension == '.md':
                chunks = markdown_splitter.split_documents([doc])
                logger.debug(f"Split {file_path} as Markdown.")
            elif file_extension in ['.txt', '.json', '.xml', '.yml', '.yaml', '.csv']: # Add other common text files
                chunks = generic_splitter.split_documents([doc])
                logger.debug(f"Split {file_path} as generic text.")
            else:
                # Skip binary files or unhandled types
                continue
            
            processed_docs.extend(chunks)

        return processed_docs

    except Exception as e:
        logger.error(f"Error cloning or processing GitHub repo {repo_url}: {e}")
        return []
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            logger.info(f"Cleaned up temporary directory: {temp_dir}")

# --- Interview Logic Functions ---

async def _get_relevant_context(session_data: Dict[str, Any], query: str) -> str:
    """Retrieves relevant chunks from the ChromaDB vector store."""
    chroma_db_path = session_data.get('chroma_db_path')
    if not chroma_db_path:
        logger.warning("ChromaDB path not found in session data.")
        return ""
    try:
        vector_store = Chroma(persist_directory=chroma_db_path, embedding_function=embeddings)
        retrieved_docs = vector_store.similarity_search(query, k=3)
        context = "\n---\n".join([doc.page_content for doc in retrieved_docs])
        logger.info(f"Retrieved {len(retrieved_docs)} documents from ChromaDB for query: '{query[:50]}...'")
        return context
    except Exception as e:
        logger.error(f"Error retrieving from ChromaDB at {chroma_db_path} for query '{query[:50]}...': {e}")
        return ""

async def generate_question(session_data: Dict[str, Any], question_type: str, user_answer_prev: str = "") -> str:
    """Generates a new interview question based on type and context."""
    session_data['current_question_count'] += 1
    prompt_template_str = "" # Renamed to avoid conflict with PromptTemplate object
    context_query = ""
    context_data = ""
    github_context = ""

    # Access session state correctly
    sections_covered = session_data['sections_covered']
    section_questions_asked = session_data['section_questions_asked']
    max_questions_per_section = session_data['max_questions_per_section']

    logger.info(f"Generating question for type: {question_type}, asked: {section_questions_asked[question_type]}")

    if question_type == 'introduction' and not sections_covered['introduction']:
        if section_questions_asked['introduction'] == 0:
            prompt_template_str = "You are an AI interviewer. Begin the mock interview by asking the candidate to introduce themselves. Be polite and professional. Question:"
        else:
            prompt_template_str = "The candidate just introduced themselves: '{user_answer_prev}'. Based on this, ask a follow-up question to probe deeper into their background or motivations. Question:"
            context_data = user_answer_prev
        section_questions_asked['introduction'] += 1
        if section_questions_asked['introduction'] >= max_questions_per_section + 1:
            sections_covered['introduction'] = True

    elif question_type == 'skills' and not sections_covered['skills']:
        context_query = "candidate's skills, technical proficiencies, software, tools, languages"
        relevant_resume_skills = await _get_relevant_context(session_data, context_query)
        if section_questions_asked['skills'] == 0:
            prompt_template_str = """You are an AI interviewer. Based on the candidate's resume, specifically their skills section, ask a question that assesses their proficiency or experience with a key skill.
            Resume Skills Context:
            {context}
            Question:"""
        else:
            prompt_template_str = """You are an AI interviewer. The candidate just answered a question about their skills. Based on their resume and previous answer, ask another question that probes deeper into a specific skill or how they applied it.
            Resume Skills Context:
            {context}
            Previous Answer: {user_answer_prev}
            Question:"""
        context_data = relevant_resume_skills
        section_questions_asked['skills'] += 1
        if section_questions_asked['skills'] >= max_questions_per_section:
            sections_covered['skills'] = True

    elif question_type == 'projects' and not sections_covered['projects']:
        context_query = "candidate's projects, personal projects, portfolio, github repositories, code examples"
        relevant_resume_projects = await _get_relevant_context(session_data, context_query)
        github_knowledge_str = "\n".join(session_data['github_project_knowledge'])
        if section_questions_asked['projects'] == 0:
            prompt_template_str = """You are an AI interviewer. Based on the candidate's resume projects and any provided GitHub project knowledge (code snippets/READMEs), ask a question about a specific project, their role, or the technologies used.
            Resume Projects Context:
            {context}
            GitHub Project Knowledge:
            {github_knowledge}
            Question:"""
        else:
            prompt_template_str = """You are an AI interviewer. The candidate just answered a question about their projects. Based on their resume, GitHub project knowledge, and previous answer, ask a follow-up question about a specific project, challenges faced, or impact achieved.
            Resume Projects Context:
            {context}
            GitHub Project Knowledge:
            {github_knowledge}
            Previous Answer: {user_answer_prev}
            Question:"""
        context_data = relevant_resume_projects
        github_context = github_knowledge_str
        section_questions_asked['projects'] += 1
        if section_questions_asked['projects'] >= max_questions_per_section:
            sections_covered['projects'] = True

    elif question_type == 'experience' and not sections_covered['experience']:
        context_query = "candidate's work experience, professional roles, job history, responsibilities, achievements"
        relevant_resume_experience = await _get_relevant_context(session_data, context_query)
        if section_questions_asked['experience'] == 0:
            prompt_template_str = """You are an AI interviewer. Based on the candidate's work experience in their resume, ask a question about a past role, responsibility, or achievement.
            Resume Experience Context:
            {context}
            Question:"""
        else:
            prompt_template_str = """You are an AI interviewer. The candidate just answered a question about their experience. Based on their resume and previous answer, ask a follow-up question to delve deeper into a specific experience, challenge, or learning.
            Resume Experience Context:
            {context}
            Previous Answer: {user_answer_prev}
            Question:"""
        context_data = relevant_resume_experience
        section_questions_asked['experience'] += 1
        if section_questions_asked['experience'] >= max_questions_per_section:
            sections_covered['experience'] = True
    else:
        logger.info("All question sections covered. Signalling INTERVIEW_FINISHED.")
        return "INTERVIEW_FINISHED"

    if not prompt_template_str:
        logger.error("No prompt template generated for current question type. Signalling INTERVIEW_FINISHED.")
        return "INTERVIEW_FINISHED"

    prompt_obj = PromptTemplate.from_template(prompt_template_str) # Renamed to prompt_obj
    
    # Prepare all possible input variables for the prompt, providing defaults if not directly used
    formatted_inputs = {
        "context": context_data,
        "github_knowledge": github_context,
        "user_answer_prev": user_answer_prev,
        "introduction": user_answer_prev # 'introduction' is only used in specific intro prompts
    }
    
    # Filter inputs to only include those actually defined in the current prompt_template
    final_inputs = {k: v for k, v in formatted_inputs.items() if "{" + k + "}" in prompt_obj.template}

    # Add format instructions for the StructuredOutputParser
    full_prompt_template_str = prompt_template_str + "\n\n{format_instructions}"
    prompt_obj_with_parser = PromptTemplate(
        template=full_prompt_template_str,
        input_variables=list(final_inputs.keys()) + ["format_instructions"], # Ensure format_instructions is an input variable
        partial_variables={"format_instructions": question_format_instructions}
    )

    prompt_messages = prompt_obj_with_parser.format_prompt(**final_inputs).to_messages()

    try:
        # Directly invoke the LLM with the formatted messages
        ai_response_obj = await llm.ainvoke(prompt_messages) # Invoke LLM, get AIMessage object
        ai_response_content = ai_response_obj.content # Extract content string
        
        # Parse the content using the StructuredOutputParser
        # This is where the LLM's raw output is expected to be a JSON string
        parsed_question = question_parser.parse(ai_response_content)
        question_text = parsed_question.question # Extract the actual question string
        
        session_data['conversation_history'].append(('ai', question_text)) # Store only the question
        logger.info(f"AI generated question: {question_text[:100]}...")
        return question_text
    except Exception as e:
        logger.error(f"Error invoking LLM for question generation or parsing: {e}")
        # If parsing fails, it means the LLM didn't return valid JSON.
        # We can try to return the raw content if it's not too long, or a generic error.
        if isinstance(e, json.JSONDecodeError):
            logger.error(f"LLM did not return valid JSON for question. Raw content: {ai_response_content[:500]}...")
            return "I'm sorry, I encountered an error generating a structured question. Please try again."
        return "I'm sorry, I encountered an error generating the question. Let's move to the next section or end the interview."

def _determine_next_question_type(session_data: Dict[str, Any]) -> str:
    """Determines the next section for questions based on coverage."""
    sections_covered = session_data['sections_covered']
    if not sections_covered['introduction']:
        return 'introduction'
    if not sections_covered['skills']:
        return 'skills'
    if not sections_covered['projects']:
        return 'projects'
    if not sections_covered['experience']:
        return 'experience'
    logger.info("Determined next stage: feedback_stage (all question sections covered).")
    return 'feedback_stage'

async def get_feedback_agents(session_data: Dict[str, Any]) -> Tuple[TechnicalFeedback, HRFeedback]:
    """Generates technical and HR feedback based on the full conversation."""
    full_conversation_str = "\n".join([f"{role.upper()}: {text}" for role, text in session_data['conversation_history']])
    logger.info("Generating feedback from technical and HR agents.")

    # Initialize response objects to None for safe access in except block
    tech_response_obj = None
    hr_response_obj = None

    # Technical Agent Prompt
    technical_parser = PydanticOutputParser(pydantic_object=TechnicalFeedback) # Use PydanticOutputParser
    technical_prompt = PromptTemplate(
        template="""You are a technical interviewer and expert. Analyze the following mock interview conversation.
        Based on the candidate's answers, especially those related to skills, projects, and experience,
        rate their technical knowledge on a scale of 1 to 5 (1=Poor, 5=Excellent).
        Then, provide 3-5 specific, actionable tips to help them improve their technical responses and knowledge.

        Conversation:
        {conversation}

        {format_instructions}
        """,
        input_variables=["conversation"],
        partial_variables={"format_instructions": technical_parser.get_format_instructions()},
    )
    technical_chain = technical_prompt | llm # Define chain here

    # HR Agent Prompt
    hr_parser = PydanticOutputParser(pydantic_object=HRFeedback) # Use PydanticOutputParser
    hr_prompt = PromptTemplate(
        template="""You are an HR interviewer and communication expert. Analyze the following mock interview conversation.
        Based on the candidate's overall communication, clarity, confidence, and how they structured their answers (e.g., STAR method for behavioral questions),
        rate their communication skills on a scale of 1 to 5 (1=Poor, 5=Excellent).
        Then, provide 3-5 specific, actionable tips to help them improve their communication and soft skills.

        Conversation:
        {conversation}

        {format_instructions}
        """,
        input_variables=["conversation"],
        partial_variables={"format_instructions": hr_parser.get_format_instructions()},
    )
    hr_chain = hr_prompt | llm # Define chain here

    try:
        # Invoke the chains and get the raw AIMessage objects
        tech_response_obj = await technical_chain.ainvoke({"conversation": full_conversation_str})
        hr_response_obj = await hr_chain.ainvoke({"conversation": full_conversation_str})

        # Explicitly parse the content using the respective parsers
        technical_feedback = technical_parser.parse(tech_response_obj.content)
        hr_feedback = hr_parser.parse(hr_response_obj.content)

        logger.info(f"Technical feedback generated. Rating: {technical_feedback.technical_knowledge_rating}/5")
        logger.info(f"HR feedback generated. Rating: {hr_feedback.communication_skills_rating}/5")
        return technical_feedback, hr_feedback
    except Exception as e:
        logger.error(f"Error generating feedback: {e}")
        # Log raw LLM outputs if parsing failed to help debug
        if 'tech_response_obj' in locals() and tech_response_obj: # Check if defined
            logger.error(f"Raw LLM output for technical feedback (if available): {tech_response_obj.content[:500]}...")
        if 'hr_response_obj' in locals() and hr_response_obj: # Check if defined
            logger.error(f"Raw LLM output for HR feedback (if available): {hr_response_obj.content[:500]}...")
        
        # Return default/empty feedback to prevent a full crash
        return TechnicalFeedback(technical_knowledge_rating=0, technical_tips=["Error generating feedback."]), \
               HRFeedback(communication_skills_rating=0, communication_tips=["Error generating feedback."])

# --- FastAPI Endpoints ---

@app.post("/register")
async def register_user(username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    logger.info(f"Attempting to register user: {username}")
    user = db.query(User).filter(User.username == username).first()
    if user:
        logger.warning(f"Registration failed: Username '{username}' already exists.")
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Username already registered")
    
    hashed_password = get_password_hash(password)
    new_user = User(username=username, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    logger.info(f"User '{username}' registered successfully with ID: {new_user.id}")
    return JSONResponse(content={"message": "User registered successfully"})

@app.post("/token")
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    logger.info(f"Attempting to log in user: {form_data.username}")
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        logger.warning(f"Login failed for user '{form_data.username}': Incorrect credentials.")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token = create_access_token(data={"sub": user.username})
    logger.info(f"User '{form_data.username}' logged in successfully.")
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/upload_resume")
async def upload_resume_endpoint(
    resume_file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Receives resume file, extracts text, finds GitHub links, clones repos,
    and creates ChromaDB vector store for the session.
    """
    session_id = str(current_user.id) + "_" + datetime.now().strftime("%Y%m%d%H%M%S%f")
    logger.info(f"User {current_user.username} (ID: {current_user.id}) uploading resume. Session ID: {session_id}")

    temp_dir = tempfile.mkdtemp(prefix=f"resume_upload_{session_id}_")
    file_path = os.path.join(temp_dir, resume_file.filename)
    logger.info(f"Saving uploaded file to temporary path: {file_path}")

    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(resume_file.file, buffer)

        resume_text = ""
        file_extension = os.path.splitext(resume_file.filename)[1].lower()
        if file_extension == '.pdf':
            resume_text = extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            resume_text = extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                resume_text = f.read()
        else:
            logger.error(f"Unsupported file type uploaded: {file_extension}")
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

        if not resume_text:
            logger.error("Failed to extract text from uploaded resume.")
            raise HTTPException(status_code=500, detail="Failed to extract text from resume.")
        logger.info(f"Resume text extracted. Length: {len(resume_text)} characters.")

        chroma_db_path = os.path.join(tempfile.gettempdir(), f"chroma_db_{session_id}")
        os.makedirs(chroma_db_path, exist_ok=True)
        logger.info(f"Created ChromaDB directory: {chroma_db_path}")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        resume_chunks = text_splitter.split_text(resume_text)
        resume_docs = [Document(page_content=chunk, metadata={"source": "resume"}) for chunk in resume_chunks]
        logger.info(f"Split resume into {len(resume_chunks)} chunks.")

        Chroma.from_documents(
            documents=resume_docs,
            embedding=embeddings,
            persist_directory=chroma_db_path
        )
        logger.info(f"Resume text embedded into ChromaDB at {chroma_db_path}")

        github_links = extract_github_links(resume_text)
        github_project_knowledge = []
        for link in github_links:
            logger.info(f"Processing GitHub link: {link}")
            code_documents = await clone_and_embed_github_repo(link, session_id)
            if code_documents:
                Chroma.from_documents(
                    documents=code_documents,
                    embedding=embeddings,
                    persist_directory=chroma_db_path
                )
                logger.info(f"Embedded code from {link} into ChromaDB.")
                github_project_knowledge.append(f"Content from {link} processed.")
            else:
                logger.warning(f"No code documents found or processed from {link}.")

        # Create a new interview session record in DB
        new_interview_session = InterviewSession(
            user_id=current_user.id,
            resume_text_snippet=resume_text[:500] + "..." if len(resume_text) > 500 else resume_text,
            github_knowledge_summary=json.dumps(github_project_knowledge) if github_project_knowledge else None,
            start_time=datetime.now()
        )
        db.add(new_interview_session)
        db.commit()
        db.refresh(new_interview_session)
        logger.info(f"New interview session DB record created with ID: {new_interview_session.id}")

        # Initialize active session data
        active_sessions[session_id] = {
            "db_session_id": new_interview_session.id, # Link to DB record
            "resume_text": resume_text,
            "github_project_knowledge": github_project_knowledge,
            "chroma_db_path": chroma_db_path,
            "conversation_history": [],
            "current_question_count": 0,
            "sections_covered": {
                'introduction': False,
                'skills': False,
                'projects': False,
                'experience': False
            },
            "section_questions_asked": {
                'introduction': 0,
                'skills': 0,
                'projects': 0,
                'experience': 0
            },
            "max_questions_per_section": 2
        }
        logger.info(f"Active session data initialized for session ID: {session_id}")

        return JSONResponse(content={"message": "Resume processed and embeddings created.", "interview_session_id": session_id})

    except Exception as e:
        logger.exception(f"Critical error during resume upload for user {current_user.username}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process resume: {str(e)}")
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            logger.info(f"Cleaned up temporary resume directory: {temp_dir}")


@app.post("/start_interview")
async def start_interview_endpoint(
    interview_session_id: str = Form(...),
    current_user: User = Depends(get_current_user) # Ensure user is authenticated
):
    """Starts the mock interview for a given session."""
    logger.info(f"User {current_user.username} starting interview for session ID: {interview_session_id}")
    session_data = active_sessions.get(interview_session_id)
    if not session_data:
        logger.error(f"Attempted to start interview for non-existent active session: {interview_session_id}")
        raise HTTPException(status_code=404, detail="Active interview session not found. Please upload resume first.")

    # Ensure this session belongs to the current user (basic check)
    if not interview_session_id.startswith(str(current_user.id) + "_"):
        logger.error(f"Unauthorized access attempt to session {interview_session_id} by user {current_user.username}")
        raise HTTPException(status_code=403, detail="Unauthorized access to session.")


    # Reset interview state for a fresh start if already exists in active_sessions
    session_data["conversation_history"] = []
    session_data["current_question_count"] = 0
    session_data["sections_covered"] = {
        'introduction': False,
        'skills': False,
        'projects': False,
        'experience': False
    }
    session_data["section_questions_asked"] = {
        'introduction': 0,
        'skills': 0,
        'projects': 0,
        'experience': 0
    }
    logger.info(f"Interview state reset for session ID: {interview_session_id}")

    try:
        initial_question = await generate_question(session_data, 'introduction')
        logger.info(f"Initial question generated for session {interview_session_id}.")
        return JSONResponse(content={"question": initial_question})
    except Exception as e:
        logger.exception(f"Error starting interview for session {interview_session_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to start interview: {str(e)}")

@app.post("/submit_answer")
async def submit_answer_endpoint(
    interview_session_id: str = Form(...),
    user_answer: str = Form(...),
    current_user: User = Depends(get_current_user), # Ensure user is authenticated
    db: Session = Depends(get_db)
):
    """Submits user's answer and generates the next question."""
    logger.info(f"User {current_user.username} submitting answer for session {interview_session_id}.")
    session_data = active_sessions.get(interview_session_id)
    if not session_data:
        logger.error(f"Attempted to submit answer for non-existent active session: {interview_session_id}")
        raise HTTPException(status_code=404, detail="Active interview session not found. Please upload resume first.")

    if not user_answer.strip():
        logger.warning(f"Empty answer submitted for session {interview_session_id}.")
        raise HTTPException(status_code=400, detail="User answer cannot be empty.")

    # Save user answer to DB
    new_conversation_entry = InterviewConversation(
        interview_session_id=session_data["db_session_id"],
        role='user',
        text=user_answer,
        timestamp=datetime.now()
    )
    db.add(new_conversation_entry)
    db.commit() # Commit immediately to ensure it's saved
    logger.info(f"User answer saved to DB for session {interview_session_id}.")

    session_data['conversation_history'].append(('user', user_answer))

    next_question_type = _determine_next_question_type(session_data)
    if next_question_type == 'feedback_stage':
        logger.info(f"Interview questions completed for session {interview_session_id}. Moving to feedback stage.")
        return JSONResponse(content={"status": "interview_finished"})

    try:
        next_ai_question = await generate_question(session_data, next_question_type, user_answer)

        # Save AI question to DB
        new_conversation_entry = InterviewConversation(
            interview_session_id=session_data["db_session_id"],
            role='ai',
            text=next_ai_question,
            timestamp=datetime.now()
        )
        db.add(new_conversation_entry)
        db.commit() # Commit immediately
        logger.info(f"AI question saved to DB for session {interview_session_id}.")

        if next_ai_question == "INTERVIEW_FINISHED":
             logger.info(f"AI explicitly signalled INTERVIEW_FINISHED for session {interview_session_id}.")
             return JSONResponse(content={"status": "interview_finished"})
        return JSONResponse(content={"question": next_ai_question})
    except Exception as e:
        logger.exception(f"Error submitting answer or generating next question for session {interview_session_id}: {e}")
        db.rollback() # Rollback if error
        raise HTTPException(status_code=500, detail=f"Failed to generate next question: {str(e)}")

@app.post("/get_feedback")
async def get_interview_feedback_endpoint(
    interview_session_id: str = Form(...),
    current_user: User = Depends(get_current_user), # Ensure user is authenticated
    db: Session = Depends(get_db)
):
    """Generates and returns technical and HR feedback."""
    logger.info(f"User {current_user.username} requesting feedback for session {interview_session_id}.")
    session_data = active_sessions.get(interview_session_id)
    if not session_data:
        logger.error(f"Attempted to get feedback for non-existent active session: {interview_session_id}")
        raise HTTPException(status_code=404, detail="Active interview session not found. Please upload resume first.")

    # Define parsers and prompts within the function scope
    # Technical Agent Prompt
    technical_parser = PydanticOutputParser(pydantic_object=TechnicalFeedback)
    technical_prompt = PromptTemplate(
        template="""You are a technical interviewer and expert. Analyze the following mock interview conversation.
        Based on the candidate's answers, especially those related to skills, projects, and experience,
        rate their technical knowledge on a scale of 1 to 5 (1=Poor, 5=Excellent).
        Then, provide 3-5 specific, actionable tips to help them improve their technical responses and knowledge.

        Conversation:
        {conversation}

        {format_instructions}
        """,
        input_variables=["conversation"],
        partial_variables={"format_instructions": technical_parser.get_format_instructions()},
    )
    technical_chain = technical_prompt | llm

    # HR Agent Prompt
    hr_parser = PydanticOutputParser(pydantic_object=HRFeedback)
    hr_prompt = PromptTemplate(
        template="""You are an HR interviewer and communication expert. Analyze the following mock interview conversation.
        Based on the candidate's overall communication, clarity, confidence, and how they structured their answers (e.g., STAR method for behavioral questions),
        rate their communication skills on a scale of 1 to 5 (1=Poor, 5=Excellent).
        Then, provide 3-5 specific, actionable tips to help them improve their communication and soft skills.

        Conversation:
        {conversation}

        {format_instructions}
        """,
        input_variables=["conversation"],
        partial_variables={"format_instructions": hr_parser.get_format_instructions()},
    )
    hr_chain = hr_prompt | llm

    full_conversation_str = "\n".join([f"{role.upper()}: {text}" for role, text in session_data['conversation_history']])
    logger.info("Generating feedback from technical and HR agents.")

    # Initialize response objects to None for safe access in except block
    tech_response_obj = None
    hr_response_obj = None

    try:
        # Invoke the chains and get the raw AIMessage objects
        tech_response_obj = await technical_chain.ainvoke({"conversation": full_conversation_str})
        hr_response_obj = await hr_chain.ainvoke({"conversation": full_conversation_str})

        # Explicitly parse the content using the respective parsers
        technical_feedback = technical_parser.parse(tech_response_obj.content)
        hr_feedback = hr_parser.parse(hr_response_obj.content)

        logger.info(f"Technical feedback generated. Rating: {technical_feedback.technical_knowledge_rating}/5")
        logger.info(f"HR feedback generated. Rating: {hr_feedback.communication_skills_rating}/5")
        
        # Save feedback to DB
        new_feedback_entry = InterviewFeedback(
            interview_session_id=session_data["db_session_id"],
            technical_rating=technical_feedback.technical_knowledge_rating,
            technical_tips=json.dumps(technical_feedback.technical_tips), # Store as JSON string
            hr_rating=hr_feedback.communication_skills_rating,
            hr_tips=json.dumps(hr_feedback.communication_tips) # Store as JSON string
        )
        db.add(new_feedback_entry)

        # Update interview session end time
        interview_session_db = db.query(InterviewSession).filter(InterviewSession.id == session_data["db_session_id"]).first()
        if interview_session_db:
            interview_session_db.end_time = datetime.now()
        db.commit()
        logger.info(f"Feedback and session end time saved to DB for session {interview_session_id}.")

        # Clean up ChromaDB directory after feedback is generated
        if os.path.exists(session_data['chroma_db_path']):
            shutil.rmtree(session_data['chroma_db_path'])
            logger.info(f"Cleaned up ChromaDB directory: {session_data['chroma_db_path']}")
        del active_sessions[interview_session_id] # Remove session from memory
        logger.info(f"Session {interview_session_id} removed from active sessions.")

        return JSONResponse(content={
            "technical_feedback": technical_feedback.dict(),
            "hr_feedback": hr_feedback.dict()
        })
    except Exception as e:
        logger.exception(f"Error getting feedback for session {interview_session_id}: {e}")
        # Log raw LLM outputs if parsing failed to help debug
        if 'tech_response_obj' in locals() and tech_response_obj:
            logger.error(f"Raw LLM output for technical feedback (if available): {tech_response_obj.content[:500]}...")
        if 'hr_response_obj' in locals() and hr_response_obj:
            logger.error(f"Raw LLM output for HR feedback (if available): {hr_response_obj.content[:500]}...")
        
        db.rollback() # Rollback if error
        raise HTTPException(status_code=500, detail=f"Failed to generate feedback: {str(e)}")

@app.get("/history")
async def get_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """Retrieves interview history for the current user."""
    logger.info(f"User {current_user.username} requesting interview history.")
    history_records = db.query(InterviewSession).filter(InterviewSession.user_id == current_user.id).all()

    response_data = []
    for session in history_records:
        conversations = []
        # Order conversations by timestamp to ensure correct sequence
        for conv in sorted(session.conversations, key=lambda c: c.timestamp):
            conversations.append({"role": conv.role, "text": conv.text})

        feedback_data = None
        hr_feedback_data = None
        if session.feedback:
            feedback_data = {
                "rating": session.feedback.technical_rating,
                "tips": json.loads(session.feedback.technical_tips) if session.feedback.technical_tips else []
            }
            hr_feedback_data = {
                "rating": session.feedback.hr_rating,
                "tips": json.loads(session.feedback.hr_tips) if session.feedback.hr_tips else []
            }

        response_data.append({
            "interview_session_id": session.id, # Use DB ID for history
            "start_time": session.start_time.isoformat(),
            "end_time": session.end_time.isoformat() if session.end_time else None,
            "resume_snippet": session.resume_text_snippet,
            "github_summary": json.loads(session.github_knowledge_summary) if session.github_knowledge_summary else [],
            "conversation": conversations,
            "technical_feedback": feedback_data,
            "hr_feedback": hr_feedback_data
        })
    # Sort by start_time descending
    response_data.sort(key=lambda x: x['start_time'], reverse=True)
    logger.info(f"Retrieved {len(response_data)} history records for user {current_user.username}.")
    return JSONResponse(content=response_data)

@app.get("/")
async def serve_frontend():
    """Serves the main HTML frontend file."""
    frontend_path = os.path.join(os.getcwd(), "static", "index.html")
    logger.info(f"Attempting to serve frontend from: {frontend_path}")
    if not os.path.exists(frontend_path):
        logger.error(f"Frontend file not found at: {frontend_path}")
        raise HTTPException(status_code=404, detail="Frontend file not found.")
    return FileResponse(frontend_path)
